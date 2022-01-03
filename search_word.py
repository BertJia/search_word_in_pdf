# 2021-4-29 19:00:12
# 花了两天的时间用Python实现了索引的建立， 能够替代word中的索引+excel的处理
# 最开始想的是在word中搜索对应的页码，结果网上说word是流式的，不打开根本不知道页码，意外发现可以使用PDF，
# 在pdf文件中搜索关键词所在的页码，想这还简单了，本来还需要PDF转成word才能做，这次省了。
# 目前搜索到的词不会漏，但可能会多，因为读取pdf中页脚的内容也算到正文里的，如果在页脚内有关键词，就会大量的出现
# 2022-01-03 16:41:10
# 1. 在*.docx 结尾的word文档中指定关键词或词组，每个占一行



import os
import datetime as datetime

from docx import Document
import fitz   # 需要安装扩展库pymupdf


# 从word文档中把要索引的条目取出来
# 去除了空行， 条目两侧的空格
def get_word_item(file_path):
    doc = Document(file_path)
    word_items = []
    for paragraph in doc.paragraphs:
        # print(paragraph.text)
        if paragraph.text.strip() != '':
            word_items.append(paragraph.text.strip())
    return word_items


# 对条目进行加工
# 每个条目组成一个list， 第一个为原始词，后面包含 条目可能换行，需要增加 a\nb, 可能条目后带有括号，需要把括号中内容去掉
def process_word_item(word_items, add_brk=False):
    processed_word_list = []
    for word in word_items:
        # 此次进行加换行符， 去括号的处理
        # 此次想第一次执行时，就把索引条目打散，加换行符，这样会使索引的关键词成倍的增加，所以考虑先进行一次筛选
        # 然后把未找到的再进行一次循环，但是这样存在两个问题
        # 1. 可能找到的没有找全，会漏掉恰好也换行的那个
        # 2. 再次找到的会排在后面，怎么在找到的序列中按最开始的索引项排序, (这个已解决，在后面进行了排序)

        if word.find("（") != -1:
            s_word = word.split("（")
            e_word = word.split("）")
            f_word = s_word[0] + e_word[len(e_word)-1]
        else:
            f_word = word
        if add_brk:
            # 增加换行符， 换页符
            p_word = [word]  # 第一个为原始词，用在最后输出时, 从第二个开始在文档中查找
            for i, w in enumerate(f_word, start=1):
                if i < len(f_word):
                    # 换行
                    ret_word = f_word.replace(w, w+'\n', 1)
                    p_word.append(ret_word)
                    # 换页
                    # 存在问题，一个词语只判断第一个字最本页的最后，不能确定就是这个词。
                    # p_word.append('\n'+f_word[-1*i:])  # 下一页开头的词,这种会把索引页码显示为下一页
                    p_word.append(f_word[0:i]+'\n')  # 本页末尾的词
            print(p_word)
        else:
            p_word = [word, f_word]  # 第一个为原始词，用在最后输出时, 从第二个开始在文档中查找
        processed_word_list.append(p_word)
    return processed_word_list


# 从pdf文件中查找关键词所在的页码，需要指定，从第几页开始到第几页结束
def find_pages_by_word(pdf_file, keyword, s_page=1, e_page=10):
    pages = []
    # 逐页提取并检查PDF文件中的文本
    with fitz.open(pdf_file) as doc:
        for index, page in enumerate(doc, start=1):
            # 索引从正文开始，掐头去尾
            if s_page <= index <= e_page:
                curr_page_txt = page.getText()
                if curr_page_txt != '':
                    # 此处是为了处理字符跨页
                    if keyword.find('\n') and keyword.rfind('\n') == len(keyword) - 1:
                        # 判断是最后一个
                        if curr_page_txt.find(keyword) and \
                                curr_page_txt.rfind(keyword) == len(curr_page_txt) - len(keyword):
                            pages.append(index - s_page + 1)
                    else:
                        if keyword in curr_page_txt:
                            # print(f'第{index}页包含关键字')
                            # 索引需要从正文开始算第一页, 此次做减法，使关键词在正文的页码上
                            pages.append(index - s_page + 1)
        return pages


letter_list = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M',
               'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']


def gen_indexes(word_items, book_pdf_file, s_page, e_page):
    ret_obj = {'find': {}, 'unfounded': []}
    for words in word_items:
        ret_pages = []
        for word in words[1:]:  # 从第二项开始，第一项为原始词
            # 如果是字母A-Z不进行查找
            if word in letter_list:
                continue
            inner_pages = find_pages_by_word(book_pdf_file, word, s_page, e_page)
            ret_pages.append(inner_pages)

        tmp_list = []
        for i in ret_pages:
            for j in i:
                # tmp_list.append(str(j))
                tmp_list.append(j)
        # 此次进行排序, 当一个词分多次查找时，返回值需要进行排序
        tmp_list.sort()
        # 转为字符串, 拼接间隔符
        tmp_list_str = [str(i) for i in tmp_list]
        if len(tmp_list) > 0 or words[0] in letter_list:
            print(words[0] + '    ' + ', '.join(tmp_list_str))
            ret_obj['find'].setdefault(words[0], ', '.join(tmp_list_str))
        else:
            print(words[0])
            ret_obj['unfounded'].append(words[0])
    return ret_obj


def write_indexes_file(new_doc, ret_indexes, ret_type):
    file_path = new_doc
    doc = Document(file_path)
    # 这里是因为ret_indexes 可能为字典，也可能为列表，使用ret_type进行区分，有好办法时再修改
    if ret_type == 1:
        for key, item in ret_indexes.items():
            display_txt = key+'    '+item
            doc.add_paragraph(display_txt)
    elif ret_type == 2:
        for item in ret_indexes:
            display_txt = item
            doc.add_paragraph(display_txt)
    # 保存
    doc.save(file_path)


# 字典key 根据 列表排序, 自己先写循环，有好方法时再替换
def sort_keyword_by_rule(init_dict, init_list):
    print('init_dict--------', init_dict)
    print('init_list--------', init_list)
    # 需要把rule 组成一个字典
    new_dict = {}
    if len(init_list) > 0:
        for item in init_list:
            if len(init_dict) > 0:
                for d_k, d_item in list(init_dict.items()):  # 转为列表是为了删除元素
                    if item == d_k:
                        new_dict.setdefault(d_k, d_item)
                        init_dict.pop(d_k)
                        break
    return new_dict


# 开始执行

os.chdir(r"E:\doc")
# 查看当前工作目录
curr_dir = os.getcwd()
print("当前工作目录为 %s" % curr_dir)


indexes_path = '关键词.docx'
book_pdf_path = 'Python语法整理.pdf'
start_page = 1  # 正文开始页码
end_page = 60  # 正文结束页码


# test start
# start_page = 27
# end_page = 178
# test end

# start_time = time.strftime("%Y-%m-%d %H:%M:%S.%f")
start_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')
print('start_time:--------------'+start_time)
# 获取关键词
word_list_items = get_word_item(indexes_path)
# 加工关键词
# 从第二项开始，因为第一项一般都是标题 例如 “索引”, 不需要
processed_word_items = process_word_item(word_list_items[1:])

# test start
# word_list_items.insert(1, '循环')
# processed_word_items = process_word_item(word_list_items[1:50])
# test end

# search1_start_time = time.strftime("%Y-%m-%d %H:%M:%S.%f")
search1_start_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')
print('search1_start_time:--------------'+search1_start_time)
# 获取关键词索引
result_obj = gen_indexes(processed_word_items, book_pdf_path, start_page, end_page)
founded_item_indexes = result_obj['find']
unfounded_item_indexes = result_obj['unfounded']

# search1_end_time = time.strftime("%Y-%m-%d %H:%M:%S.%f")
search1_end_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')
print('search1_end_time:--------------'+search1_end_time)

if len(result_obj['unfounded']) > 0:
    print('--------没有找到的再次查找--------')

    unfounded_word_items = process_word_item(result_obj['unfounded'], True)

    # search2_start_time = time.strftime("%Y-%m-%d %H:%M:%S.%f")
    search2_start_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')
    print('search2_start_time:--------------' + search2_start_time)

    result_sec_obj = gen_indexes(unfounded_word_items, book_pdf_path, start_page, end_page)
    if len(result_sec_obj['find']) > 0:
        founded_item_indexes.update(result_sec_obj['find'])
    unfounded_item_indexes = result_sec_obj['unfounded']

    # search2_end_time = time.strftime("%Y-%m-%d %H:%M:%S.%f")
    search2_end_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')
    print('search2_end_time:--------------' + search2_end_time)

    # 此处存在一个问题，第二次查找到的关键词会排在后面。
    # 需要对result_obj['find']和result_sec_obj['find']进行排序，但是直接sort会使 英文字母都排在开始的位置
    # 思路1：在处理时再关键词前增加数字索引，这个索引一直带到最后排序后再删除。但是这个会增加程序复杂度
    # 思路2：找到一个按照指定顺序排序的方法，使结果按照 word_list_items 进行排序, 遇到的问题是 找到的结果后面拼接上了页码,
    #       和原始关键词不匹配了
    sort_start_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')
    print('sort_start_time:--------------' + sort_start_time)
    founded_item_indexes = sort_keyword_by_rule(founded_item_indexes, word_list_items)
    sort_end_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')
    print('sort_end_time:--------------' + sort_end_time)

# 写入文件
need_write_file = True
# need_write_file = False # test
if need_write_file:
    # write_start_time = time.strftime("%Y-%m-%d %H:%M:%S.%f")
    write_start_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')
    print('write_start_time:--------------' + write_start_time)
    new_file = '关键词页码-结果.docx'
    write_indexes_file(new_file, founded_item_indexes, 1)

    # new_file = '《文人画概论》索引.docx'
    new_file = '关键词页码-结果.docx'
    write_indexes_file(new_file, unfounded_item_indexes, 2)

    # write_end_time = time.strftime("%Y-%m-%d %H:%M:%S.%f")
    write_end_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')
    print('write_end_time:--------------' + write_end_time)
